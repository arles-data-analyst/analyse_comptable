# ================================================================
# Nettoyage et Analyse Comptable Automatisée
# Auteur : Arlès FANAMPINDRAINY
# ================================================================

# --------------------------
# 1. IMPORTATION DES BIBLIOTHÈQUES
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import calendar
from sklearn.linear_model import LinearRegression
from prophet import Prophet
import xlsxwriter
from matplotlib.backends.backend_pdf import PdfPages
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH




print("Bibliothèques importées.")

# --------------------------
# 2. CHARGEMENT DU FICHIER BRUT
df = pd.read_excel("factures_comptables_brutes.xlsx")
print("Fichier brut chargé.")

# --------------------------
# 3. NETTOYAGE DES DONNÉES

print("Nettoyage en cours...")

# 3.1 Renommer les colonnes
df.columns = (
    df.columns.str.strip()
              .str.lower()
              .str.replace(" ", "_")
)

# 3.2 Conversion des types
df["montant"] = pd.to_numeric(df["montant"], errors="coerce")
df["date"] = pd.to_datetime(df["date"], errors="coerce")

# 3.3 Suppression des lignes avec montants ou dates manquants
df.dropna(subset=["montant", "date"], inplace=True)

# 3.4 Colonne type : Entrée ou Sortie
df["type"] = df["montant"].apply(lambda x: "Entrée" if x >= 0 else "Sortie")

# 3.5 Extraire le mois
df["mois"] = df["date"].dt.month

# 3.6 Solde cumulé
df["solde_cumulé"] = df["montant"].cumsum()

# 3.7 Nettoyage texte
if "libellé" in df.columns:
    df["libellé"] = df["libellé"].astype(str).str.strip().str.capitalize()
if "compte" in df.columns:
    df["compte"] = df["compte"].astype(str).str.strip()

# 3.8 Export du fichier nettoyé
df.to_excel("factures_comptables_nettoyees.xlsx", index=False)
print("Données nettoyées et exportées dans factures_comptables_nettoyees.xlsx.")

# --------------------------
# 4. ANALYSE COMPTABLE

print("Analyse en cours...")

montant_total = df["montant"].sum()
total_entrees = df[df["type"] == "Entrée"]["montant"].sum()
total_sorties = df[df["type"] == "Sortie"]["montant"].sum()
resultat_par_mois = df.groupby("mois")["montant"].sum()
total_par_compte = df.groupby("compte")["montant"].sum()
top_comptes = df["compte"].value_counts().head(5)
libelles_freq = df["libellé"].value_counts().head(5)

print(f"Montant total : {montant_total:.2f} €")
print(f"Total entrées : {total_entrees:.2f} €")
print(f"Total sorties : {total_sorties:.2f} €")
print(f"Solde net : {montant_total:.2f} €")

# --------------------------
# 5. GRAPHIQUES

print("Génération des graphiques...")

# 5.1 Montant total par mois
mois_complets = pd.Series(range(1, 13), name="mois")
df_mensuel = df.groupby("mois")["montant"].sum().reindex(mois_complets).fillna(0).reset_index()

plt.figure(figsize=(10, 6))
sns.barplot(data=df_mensuel, x="mois", y="montant", color="skyblue")
plt.title("Montant total par mois")
plt.xlabel("Mois")
plt.ylabel("Montant (€)")
plt.xticks(ticks=range(12), labels=["Janv", "Fév", "Mars", "Avr", "Mai", "Juin", "Juil", "Août", "Sept", "Oct", "Nov", "Déc"])
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.tight_layout()
plt.savefig("montant_par_mois.png")
#plt.show()

# 5.2 Camembert Entrées vs Sorties
type_totaux = df.groupby("type")["montant"].apply(lambda x: abs(x).sum())

plt.figure(figsize=(6, 6))
plt.pie(type_totaux, labels=type_totaux.index, autopct='%1.1f%%', startangle=90, colors=["#66bb6a", "#ef5350"])
plt.title("Répartition des montants : Entrées vs Sorties")
plt.axis("equal")
plt.tight_layout()
plt.savefig("repartition_type_operations.png")
#plt.show()

# 5.3 Top 5 comptes
top_comptes_montants = df.groupby("compte")["montant"].sum().round(2).sort_values(ascending=False).head(5)

plt.figure(figsize=(8, 5))
sns.barplot(x=top_comptes_montants.values, y=top_comptes_montants.index, color="steelblue")
plt.title("Montants par comptes comptables (Top 5)")
plt.xlabel("Montant (€)")
plt.ylabel("Compte comptable")
plt.grid(axis='x', linestyle='--', alpha=0.5)
plt.tight_layout()
plt.savefig("top_comptes_comptables.png")
#plt.show()

# 5.4 Évolution du solde cumulé dans le temps
df = df.sort_values("date")
plt.figure(figsize=(10, 5))
sns.lineplot(data=df, x="date", y="solde_cumulé", marker="o", color="#0277bd")
plt.title("Évolution du solde cumulé dans le temps")
plt.xlabel("Date")
plt.ylabel("Solde cumulé (€)")
plt.grid(True, linestyle="--", alpha=0.5)
plt.tight_layout()
plt.savefig("evolution_solde_cumule.png")
#plt.show()

# 5.5 Comparaison Entrées / Sorties par mois
flux_par_mois = df.groupby(["mois", "type"])["montant"].sum().reset_index()
flux_par_mois["montant"] = flux_par_mois.apply(
    lambda row: abs(row["montant"]) if row["type"] == "Sortie" else row["montant"],
    axis=1
)

plt.figure(figsize=(12, 6))
sns.barplot(data=flux_par_mois, x="mois", y="montant", hue="type", palette={"Entrée": "#66bb6a", "Sortie": "#ef5350"})
plt.title("Comparaison des flux financiers par mois")
plt.xlabel("Mois")
plt.ylabel("Montant (€)")
plt.grid(axis="y", linestyle="--", alpha=0.5)
plt.tight_layout()
plt.savefig("flux_entrées_sorties_par_mois.png")
#plt.show()

# 5.6 Top 7 postes de dépenses par libellé
df_sorties = df[df["type"] == "Sortie"]
depenses_par_libelle = df_sorties.groupby("libellé")["montant"].sum().abs().sort_values(ascending=False)
top_libelles = depenses_par_libelle.head(7)

plt.figure(figsize=(10, 6))
sns.barplot(x=top_libelles.values, y=top_libelles.index, color="firebrick")
plt.title("Top 7 postes de dépenses par libellé")
plt.xlabel("Montant (€)")
plt.ylabel("Libellé")
plt.grid(axis="x", linestyle="--", alpha=0.5)
plt.tight_layout()
plt.savefig("top_depenses_libelles.png")
#plt.show()

# 5.7 Évolution mensuelle du solde cumulé
solde_mensuel = df.sort_values("date").groupby("mois")["solde_cumulé"].last().reset_index()
solde_mensuel["mois"] = solde_mensuel["mois"].apply(lambda x: calendar.month_name[x].capitalize())
ordre_mois = [calendar.month_name[i].capitalize() for i in range(1, 13)]
solde_mensuel["mois"] = pd.Categorical(solde_mensuel["mois"], categories=ordre_mois, ordered=True)
solde_mensuel = solde_mensuel.sort_values("mois")

plt.figure(figsize=(10, 5))
sns.lineplot(data=solde_mensuel, x="mois", y="solde_cumulé", marker="o", linewidth=2.5, color="orange")
plt.title("Évolution mensuelle du solde cumulé")
plt.xlabel("Mois")
plt.ylabel("Solde cumulé (€)")
plt.grid(axis='y', linestyle='--', alpha=0.5)
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig("evolution_solde_mensuel.png")
#plt.show()
#----------------------------

# 6. RATIO
# 6.1 Ratio Entrées / Sorties
# Éviter la division par zéro
if abs(total_sorties) > 0:
    ratio_entrees_sorties = total_entrees / abs(total_sorties)
    print(f"➡️ Ratio Entrées / Sorties : {ratio_entrees_sorties:.2f}")
else:
    print("❗ Pas de sorties détectées, ratio Entrées / Sorties non calculable.")

# 6.2 Moyennes mensuelles (Entrées et Sorties)
moyennes_mensuelles = df.groupby(["mois", "type"])["montant"].mean().unstack()
print("\n Moyennes mensuelles Entrées / Sorties :")
print(moyennes_mensuelles.round(2))

#6.3 Écart-type des flux mensuels
ecart_type_mensuel = df.groupby("mois")["montant"].std()
print("\n Écart-type des flux mensuels :")
print(ecart_type_mensuel.round(2))

# 6.4 Ratio solde net / total entrées (marge brute approximative)
if total_entrees > 0:
    ratio_solde_sur_entrees = montant_total / total_entrees
    print(f" Ratio Solde Net / Total Entrées : {ratio_solde_sur_entrees:.2f}")
else:
    print("❗ Pas d'entrée détectée, ratio Solde / Entrées non calculable.")

# 7. COMPARAISON
# 7.1 Comparaison Mois à Mois (Évolution du Résultat Net)
# Résultat net par mois
resultat_par_mois = df.groupby("mois")["montant"].sum().sort_index()

# Calcul de l'évolution mensuelle
evolution_mensuelle = resultat_par_mois.diff().fillna(0)

# Affichage
print("\n Évolution du résultat net Mois à Mois :")
print(evolution_mensuelle.round(2))

# 7.2 Variation relative (%) entre les mois
variation_relative = resultat_par_mois.pct_change().fillna(0) * 100

print("\n Variation relative (%) du résultat Mois à Mois :")
print(variation_relative.round(2))

# 7.3 Comparaison des comptes les plus actifs
# Montant total par compte (déjà calculé)
top5_montants = total_par_compte.sort_values(ascending=False).head(5)

# Fréquence des comptes
freq_comptes = df["compte"].value_counts().head(5)

print("\n Top 5 comptes par Montant :")
print(top5_montants.round(2))

print("\n Top 5 comptes par Fréquence :")
print(freq_comptes)

# 7.4 Comptes les plus variables (écart-type)
volatilite_comptes = df.groupby("compte")["montant"].std().sort_values(ascending=False).dropna().head(5)

print("\n Comptes les plus volatils (écart-type élevé) :")
print(volatilite_comptes.round(2))

# 8. PRÉVISIONS
# 8.1 Moyenne mobile du résultat mensuel
# Résultat net mensuel (déjà calculé)
resultat_par_mois = df.groupby("mois")["montant"].sum().sort_index()

# Moyenne glissante sur 3 mois
moyenne_mobile = resultat_par_mois.rolling(window=3).mean()

# Affichage
print("\nMoyenne glissante (3 mois) du résultat net :")
print(moyenne_mobile.round(2))

# 8.2 Visualisation de la tendance
plt.figure(figsize=(10, 5))
sns.lineplot(x=resultat_par_mois.index, y=resultat_par_mois.values, label="Résultat net mensuel", marker="o")
sns.lineplot(x=moyenne_mobile.index, y=moyenne_mobile.values, label="Moyenne mobile (3 mois)", linestyle="--")
plt.title("Tendance du résultat net mensuel")
plt.xlabel("Mois")
plt.ylabel("Montant (€)")
plt.grid(True, linestyle="--", alpha=0.5)
plt.xticks(ticks=range(1, 13), labels=["Janv", "Fév", "Mars", "Avr", "Mai", "Juin", "Juil", "Août", "Sept", "Oct", "Nov", "Déc"])
plt.legend()
plt.tight_layout()
plt.savefig("tendance_resultat_mensuel.png")
#plt.show()

# 8.3 Projection simple du mois suivant
projection = moyenne_mobile.dropna().iloc[-1]
print(f"\nProjection du résultat pour le mois suivant (méthode naïve) : {projection:.2f} €")

# 8.4 Régression linéaire simple
from sklearn.linear_model import LinearRegression
import numpy as np

# Préparation des données
X = resultat_par_mois.index.values.reshape(-1, 1)  # Mois (1 à 12)
y = resultat_par_mois.values                      # Résultat net

# Régression
model = LinearRegression()
model.fit(X, y)

# Prédiction pour les mois existants + 3 mois futurs
X_future = np.arange(1, 16).reshape(-1, 1)  # 12 mois + 3 mois prévus
y_pred = model.predict(X_future)

# Affichage de l'équation
print(f"\nRégression linéaire : y = {model.coef_[0]:.2f}x + {model.intercept_:.2f}")

# 8.4 Moyenne glissante pondérée (WMA)
# Poids croissants : les mois les plus récents comptent plus
poids = np.array([1, 2, 3])
wma = resultat_par_mois.rolling(window=3).apply(lambda x: np.dot(x, poids)/poids.sum(), raw=True)

print("\nMoyenne glissante pondérée (3 mois, poids 1-2-3) :")
print(wma.round(2))

# 8.5 Graphique complet avec projection
plt.figure(figsize=(12, 6))
# Résultat net réel
sns.lineplot(x=resultat_par_mois.index, y=resultat_par_mois.values, label="Résultat net", marker="o")

# Moyenne mobile classique
sns.lineplot(x=moyenne_mobile.index, y=moyenne_mobile.values, label="Moyenne mobile", linestyle="--")

# Moyenne glissante pondérée
sns.lineplot(x=wma.index, y=wma.values, label="WMA (1-2-3)", linestyle=":")

# Régression linéaire + projection
sns.lineplot(x=range(1, 16), y=y_pred, label="Régression + projection", linestyle="dashdot", color="orange")

# Mise en forme
plt.title("Résultat net mensuel avec tendances et projection")
plt.xlabel("Mois")
plt.ylabel("Montant (€)")
plt.xticks(ticks=range(1, 16), labels=[calendar.month_abbr[i] for i in range(1, 13)] + ["M+1", "M+2", "M+3"])
plt.grid(True, linestyle="--", alpha=0.4)
plt.legend()
plt.tight_layout()
plt.savefig("projection_resultat_mensuel.png")
#plt.show()

# 8.6 Prévision des flux avec Prophet
# Objectif : Projeter les flux financiers sur les 3 prochains mois avec la bibliothèque Prophet de Meta.

# Préparation des données. 
# Agrégation mensuelle des flux
# ds : la date et y : la valeur à prédire (ici, les flux financiers mensuels)
df_prophet = df.groupby("date").agg({"montant": "sum"}).reset_index()
df_prophet = df_prophet.rename(columns={"date": "ds", "montant": "y"})

# Agréger par mois (pas jour)
df_prophet = df_prophet.resample("MS", on="ds").sum().reset_index()

#Création et entraînement du modèle
# Initialiser le modèle
model = Prophet()

# Entraîner sur les données historiques
model.fit(df_prophet)

# Projection sur les 3 mois suivants
# Créer un DataFrame contenant les 3 mois à prédire
future = model.make_future_dataframe(periods=3, freq='MS')

# Générer les prévisions
forecast = model.predict(future)

# Affichage du graphique de projection
fig1 = model.plot(forecast)
plt.title("Projection des flux financiers (3 mois)")
plt.tight_layout()
plt.savefig("projection_flux_3_mois.png")
plt.show()
#-------------------


# 9. RAPPORT AUTOMATISE
# 9.1 Générer un rapport Excel
with pd.ExcelWriter("rapport_comptable.xlsx", engine="xlsxwriter") as writer:
    # 1. Données nettoyées
    df.to_excel(writer, sheet_name="Données Nettoyées", index=False)

    # 2. Statistiques
    stats_df = pd.DataFrame({
        "Indicateur": ["Montant total", "Total Entrées", "Total Sorties", "Solde Net"],
        "Valeur (€)": [montant_total, total_entrees, total_sorties, montant_total]
    })
    stats_df.to_excel(writer, sheet_name="Statistiques Globales", index=False)

    # 3. Résumé mensuel
    df_mensuel.to_excel(writer, sheet_name="Résumé Mensuel", index=False)
    flux_par_mois.to_excel(writer, sheet_name="Flux Mensuel", index=False)
    solde_mensuel.to_excel(writer, sheet_name="Solde Mensuel", index=False)

    # 4. Insertion des graphiques
    workbook = writer.book
    for idx, image in enumerate([
        "montant_par_mois.png",
        "repartition_type_operations.png",
        "top_comptes_comptables.png",
        "evolution_solde_cumule.png",
        "flux_entrées_sorties_par_mois.png",
        "top_depenses_libelles.png",
        "evolution_solde_mensuel.png",
        "projection_flux_3_mois.png"

    ]):
        worksheet = workbook.add_worksheet(f"Graph {idx+1}")
        writer.sheets[f"Graph {idx+1}"] = worksheet
        worksheet.insert_image("B2", image)

# 9.2 Générer un rapport PDF statique
with PdfPages("rapport_comptable.pdf") as pdf:
    for image in [
        "montant_par_mois.png",
        "repartition_type_operations.png",
        "top_comptes_comptables.png",
        "evolution_solde_cumule.png",
        "flux_entrées_sorties_par_mois.png",
        "top_depenses_libelles.png",
        "evolution_solde_mensuel.png",
        "projection_flux_3_mois.png"  # ✅ ici
    ]:
        img = plt.imread(image)
        plt.figure(figsize=(11.7, 8.3))
        plt.imshow(img)
        plt.axis("off")
        pdf.savefig()
        plt.close()


# 9.3 Génération du rapport Word enrichi
doc = Document()
doc.add_heading("Rapport Comptable Automatisé", 0)

# ----------------------
# 1. Statistiques Globales
doc.add_heading("1. Statistiques Globales", level=1)
table = doc.add_table(rows=1, cols=2)
table.style = "Light List Accent 1"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Indicateur'
hdr_cells[1].text = 'Valeur (€)'

for ind, val in zip(["Montant total", "Total Entrées", "Total Sorties", "Solde Net"],
                    [montant_total, total_entrees, total_sorties, montant_total]):
    row_cells = table.add_row().cells
    row_cells[0].text = ind
    row_cells[1].text = f"{val:,.2f}"

doc.add_paragraph()

# ----------------------
# 2. Ratios et comparaisons
doc.add_heading("2. Ratios et Comparaisons", level=1)

if abs(total_sorties) > 0:
    doc.add_paragraph(f"Ratio Entrées / Sorties : {ratio_entrees_sorties:.2f}")
else:
    doc.add_paragraph("Pas de sorties détectées, ratio non calculable.")

if total_entrees > 0:
    doc.add_paragraph(f"Ratio Solde Net / Total Entrées : {ratio_solde_sur_entrees:.2f}")
else:
    doc.add_paragraph("Pas d'entrées détectées, ratio non calculable.")

doc.add_paragraph("\nÉvolution Mois à Mois :")
for mois, evol in evolution_mensuelle.items():
    doc.add_paragraph(f"Mois {mois} : {evol:.2f} €")

doc.add_paragraph("\nVariation relative Mois à Mois :")
for mois, var in variation_relative.items():
    doc.add_paragraph(f"Mois {mois} : {var:.2f} %")

doc.add_page_break()

# ----------------------
# 3. Graphiques
doc.add_heading("3. Graphiques", level=1)

for image in [
    "montant_par_mois.png",
    "repartition_type_operations.png",
    "top_comptes_comptables.png",
    "evolution_solde_cumule.png",
    "flux_entrées_sorties_par_mois.png",
    "top_depenses_libelles.png",
    "evolution_solde_mensuel.png",
    "projection_resultat_mensuel.png",
    "projection_flux_3_mois.png"
]:
    titre = image.replace(".png", "").replace("_", " ").capitalize()
    doc.add_paragraph(titre)
    try:
        doc.add_picture(image, width=Inches(6))
    except Exception as e:
        doc.add_paragraph(f"[Erreur : image non trouvée - {image}]")

doc.add_page_break()

# ----------------------
# 4. Modélisation et prévision
doc.add_heading("4. Modélisation & Prévisions", level=1)

doc.add_paragraph("Modèle de régression linéaire utilisé pour prédire les résultats sur 3 mois.")

try:
    coef = model.coef_[0]
    intercept = model.intercept_
    doc.add_paragraph(f"Équation estimée : y = {coef:.2f}x + {intercept:.2f}")
    if coef > 0:
        doc.add_paragraph("Interprétation : tendance globale à la hausse du résultat net mensuel.")
    else:
        doc.add_paragraph("Interprétation : tendance globale à la baisse du résultat net mensuel.")
except AttributeError:
    doc.add_paragraph("Le modèle utilisé ne permet pas d’extraire une équation linéaire.")

# Moyenne glissante pondérée
doc.add_paragraph("\nMoyenne glissante pondérée (poids 1-2-3) :")
for mois, val in wma.dropna().items():
    doc.add_paragraph(f"Mois {mois} : {val:.2f} €")

# Projection simple
doc.add_paragraph(f"\nProjection naïve du mois suivant : {projection:.2f} €")

# ----------------------
# Sauvegarde
doc.save("rapport_comptable.docx")
print(" Rapport Word généré : rapport_comptable.docx")


